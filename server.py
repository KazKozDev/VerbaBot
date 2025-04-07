import asyncio
import aiohttp
from quart import Quart, request, jsonify, send_file
from quart_cors import cors
import aiosqlite
import sqlite3
from datetime import datetime, timedelta
import traceback
import os
import mimetypes
import shutil
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import io
import aiofiles
import numpy as np
import pandas as pd
import docx
import openpyxl
import requests
import json
import logging
import re
from bs4 import BeautifulSoup
import hashlib
from duckduckgo_search import ddg
import urllib.parse
from typing import List, Dict, Any, Optional
import base64
from dotenv import load_dotenv
import time
import webbrowser
from search import search_web_with_retries, fetch_and_extract_text, perform_deep_search

# Updated Vector database and NLP imports
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables from .env file
load_dotenv()

# Constants
UPLOAD_FOLDER = 'uploads'
VECTOR_DB_FOLDER = 'vector_db'
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'mp3', 'mp4', 'wav', 'xlsx', 'xls', 'docx', 'csv', 'md'}

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(VECTOR_DB_FOLDER, exist_ok=True)
os.makedirs('db', exist_ok=True)

# Initialize Quart app
app = Quart(__name__)
app = cors(app, allow_origin="*")

# Configure app settings
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
llm_model = "gemma3:12b"
image_recognition_model = "gemma3:12b"
system_prompt = "You are VerbaBot - a helpful, concise, and informative assistant. You provide complete and accurate information on any topic. Always respond directly to the user's question without repeating it back. If you don't know something, be honest about it."

# Global variables for system components
rag_system = None
calendar_rag = None
memory_system = None

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
    response.headers.add('Access-Control-Allow-Methods', 'GET,POST,DELETE,OPTIONS,PUT')
    return response

# Initialize HuggingFace embeddings model (using a small model for efficiency)
embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

class CalendarManager:
    """Manager for calendar events"""
    
    def __init__(self):
        self.db_path = 'db/assistant.db'
        # Don't call async method in constructor
        
    async def setup_db(self):
        """Set up calendar tables if they don't exist"""
        async with aiosqlite.connect(self.db_path) as db:
            await db.execute('''
                CREATE TABLE IF NOT EXISTS calendar_events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT NOT NULL,
                    start_time TIMESTAMP NOT NULL,
                    end_time TIMESTAMP NOT NULL,
                    description TEXT,
                    location TEXT,
                    category TEXT DEFAULT 'default',
                    google_id TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            await db.execute('''
                CREATE TABLE IF NOT EXISTS google_calendar_auth (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    access_token TEXT,
                    refresh_token TEXT,
                    expiry_time TIMESTAMP,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            await db.commit()
            logger.info("Calendar database tables initialized")
    
    async def get_events(self, start_date=None, end_date=None, search_term=None):
        """Get calendar events with optional filtering"""
        try:
            async with aiosqlite.connect(self.db_path) as db:
                db.row_factory = sqlite3.Row
                query = "SELECT * FROM calendar_events WHERE 1=1"
                params = []
                
                if start_date:
                    query += " AND start_time >= ?"
                    params.append(start_date)
                    
                if end_date:
                    query += " AND start_time <= ?"
                    params.append(end_date)
                    
                if search_term:
                    query += " AND (title LIKE ? OR description LIKE ? OR location LIKE ?)"
                    params.extend([f"%{search_term}%", f"%{search_term}%", f"%{search_term}%"])
                    
                query += " ORDER BY start_time ASC"
                
                async with db.execute(query, params) as cursor:
                    events = []
                    async for row in cursor:
                        events.append(dict(row))
                    return events
                    
        except Exception as e:
            logger.error(f"Error getting calendar events: {e}")
            return []
    
    async def create_event(self, event_data):
        """Create a new calendar event"""
        try:
            logger.info(f"Creating calendar event with data: {event_data}")
            async with aiosqlite.connect(self.db_path) as db:
                cursor = await db.execute('''
                    INSERT INTO calendar_events 
                    (title, start_time, end_time, description, location, category, google_id)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
                    event_data['title'],
                    event_data['start_time'],
                    event_data['end_time'],
                    event_data.get('description', ''),
                    event_data.get('location', ''),
                    event_data.get('category', 'default'),
                    event_data.get('google_id', None)
                ))
                await db.commit()
                event_id = cursor.lastrowid
                logger.info(f"Created calendar event with ID: {event_id}")
                return event_id
        except Exception as e:
            logger.error(f"Error creating calendar event: {e}")
            logger.error(traceback.format_exc())
            return None
    
    async def update_event(self, event_id, event_data):
        """Update an existing calendar event"""
        try:
            async with aiosqlite.connect(self.db_path) as db:
                await db.execute('''
                    UPDATE calendar_events 
                    SET title = ?, start_time = ?, end_time = ?, description = ?, 
                        location = ?, category = ?
                    WHERE id = ?
                ''', (
                    event_data['title'],
                    event_data['start_time'],
                    event_data['end_time'],
                    event_data.get('description', ''),
                    event_data.get('location', ''),
                    event_data.get('category', 'default'),
                    event_id
                ))
                await db.commit()
                return True
        except Exception as e:
            logger.error(f"Error updating calendar event: {e}")
            return False
    
    async def delete_event(self, event_id):
        """Delete a calendar event"""
        try:
            async with aiosqlite.connect(self.db_path) as db:
                await db.execute('DELETE FROM calendar_events WHERE id = ?', (event_id,))
                await db.commit()
                return True
        except Exception as e:
            logger.error(f"Error deleting calendar event: {e}")
            return False
    
    async def get_event(self, event_id):
        """Get a specific calendar event by ID"""
        try:
            logger.info(f"Getting calendar event with ID: {event_id}")
            async with aiosqlite.connect(self.db_path) as db:
                db.row_factory = sqlite3.Row
                async with db.execute('SELECT * FROM calendar_events WHERE id = ?', (event_id,)) as cursor:
                    row = await cursor.fetchone()
                    if row:
                        event = dict(row)
                        logger.info(f"Retrieved event: {event}")
                        return event
                    else:
                        logger.warning(f"No event found with ID: {event_id}")
                        return None
        except Exception as e:
            logger.error(f"Error getting calendar event: {e}")
            return None
    
    async def save_google_auth(self, auth_data):
        """Save Google calendar authentication data"""
        try:
            async with aiosqlite.connect(self.db_path) as db:
                # Clear any existing tokens first
                await db.execute('DELETE FROM google_calendar_auth')
                
                cursor = await db.execute('''
                    INSERT INTO google_calendar_auth 
                    (access_token, refresh_token, expiry_time)
                    VALUES (?, ?, ?)
                ''', (
                    auth_data['access_token'],
                    auth_data['refresh_token'],
                    auth_data['expiry_time']
                ))
                await db.commit()
                return True
        except Exception as e:
            logger.error(f"Error saving Google auth data: {e}")
            return False
    
    async def get_google_auth(self):
        """Get Google calendar authentication data"""
        try:
            async with aiosqlite.connect(self.db_path) as db:
                db.row_factory = sqlite3.Row
                cursor = await db.execute('SELECT * FROM google_calendar_auth ORDER BY id DESC LIMIT 1')
                row = await cursor.fetchone()
                if row:
                    return dict(row)
                return None
        except Exception as e:
            logger.error(f"Error getting Google auth data: {e}")
            return None

# Initialize calendar manager
calendar_manager = CalendarManager()

class EnhancedRAG:
    def __init__(self, model_name="gemma3:12b"):
        self.model_name = model_name
        self.vectorstore = None
        self.document_metadata = {}  # Store metadata about documents
        self.load_or_create_vectorstore()
        # Don't use create_task in constructor as there's no running event loop yet
        logger.info(f"Using Ollama model: {model_name}")

    async def sync_metadata_with_db(self):
        """Synchronize document metadata with database records"""
        try:
            logger.info("Syncing document metadata with database")
            async with aiosqlite.connect('db/assistant.db') as db:
                # Check if the documents table exists
                async with db.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='documents'") as cursor:
                    if not await cursor.fetchone():
                        logger.info("Documents table doesn't exist, skipping sync")
                        return

                # Get all documents from the database
                async with db.execute('SELECT doc_id, chat_id, filename, file_path, added_at, chunks FROM documents') as cursor:
                    rows = await cursor.fetchall()
                    
                    # Update metadata from database records
                    for row in rows:
                        doc_id, chat_id, filename, file_path, added_at, chunks = row
                        
                        # Check if file exists
                        if not os.path.exists(file_path):
                            logger.warning(f"Document file not found: {file_path}, skipping")
                            continue
                            
                        # Get file hash
                        file_hash = self.get_file_hash(file_path)
                        
                        # Update metadata
                        self.document_metadata[doc_id] = {
                            "filename": filename,
                            "path": file_path,
                            "hash": file_hash,
                            "chunks": chunks or 0,
                            "added_at": added_at,
                            "chat_id": chat_id
                        }
                    
                    logger.info(f"Synchronized {len(rows)} documents from database")
                    
                    # Save updated metadata
                    self.save_vectorstore()
                    
        except Exception as e:
            logger.error(f"Error syncing document metadata: {e}")
            logger.error(traceback.format_exc())

    def load_or_create_vectorstore(self):
        """Load existing vectorstore if available, otherwise create a new one"""
        try:
            if os.path.exists(f"{VECTOR_DB_FOLDER}/index.faiss"):
                self.vectorstore = FAISS.load_local(VECTOR_DB_FOLDER, embeddings_model, allow_dangerous_deserialization=True)
                logger.info("Loaded existing vector database")
                
                # Load document metadata if available
                if os.path.exists(f"{VECTOR_DB_FOLDER}/metadata.json"):
                    with open(f"{VECTOR_DB_FOLDER}/metadata.json", 'r') as f:
                        self.document_metadata = json.load(f)
            else:
                # Create empty vector store
                self.vectorstore = FAISS.from_documents(
                    [Document(page_content="Initialization document", metadata={"source": "init"})], 
                    embeddings_model
                )
                logger.info("Created new vector database")
        except Exception as e:
            logger.error(f"Error loading vector database: {e}")
            # Create a new one if loading fails
            self.vectorstore = FAISS.from_documents(
                [Document(page_content="Initialization document", metadata={"source": "init"})], 
                embeddings_model
            )

    def save_vectorstore(self):
        """Save the current vectorstore to disk"""
        try:
            self.vectorstore.save_local(VECTOR_DB_FOLDER)
            
            # Save document metadata
            with open(f"{VECTOR_DB_FOLDER}/metadata.json", 'w') as f:
                json.dump(self.document_metadata, f)
                
            logger.info("Vector database saved successfully")
        except Exception as e:
            logger.error(f"Error saving vector database: {e}")

    def get_file_hash(self, file_path):
        """Generate a hash for a file to track changes"""
        try:
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
            return file_hash
        except Exception as e:
            logger.error(f"Error generating file hash: {e}")
            return None

    def get_text_splitter(self, file_type):
        """Return appropriate text splitter based on file type"""
        if file_type == "markdown":
            return MarkdownTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        else:
            return RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE, 
                chunk_overlap=CHUNK_OVERLAP,
                separators=["\n\n", "\n", ". ", " ", ""]
            )

    async def process_document(self, file_path, chat_id=None):
        """Process a document, split it into chunks, and add to vectorstore"""
        try:
            # Extract text content from file
            file_extension = os.path.splitext(file_path)[1].lower()
            content = await self.extract_text_from_file(file_path)
            if not content:
                return False, "Could not extract content from file"

            # Generate file identifier
            file_name = os.path.basename(file_path)
            file_hash = self.get_file_hash(file_path)
            doc_id = f"{file_name}_{file_hash}"
            
            # Check if we've already processed this exact file
            if doc_id in self.document_metadata:
                logger.info(f"Document {file_name} already processed, skipping")
                return True, "Document already processed"
            
            # Determine appropriate splitter
            splitter = self.get_text_splitter("markdown" if file_extension == ".md" else "text")
            
            # Split the text into chunks
            chunks = splitter.split_text(content)
            logger.info(f"Split document into {len(chunks)} chunks")
            
            # Create Document objects with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                if len(chunk.strip()) > 10:  # Skip very small chunks
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "filename": file_name,
                            "chat_id": chat_id,
                            "chunk_id": i,
                            "doc_id": doc_id
                        }
                    )
                    documents.append(doc)
            
            # Add to vectorstore
            self.vectorstore.add_documents(documents)
            
            # Store metadata about this document
            self.document_metadata[doc_id] = {
                "filename": file_name,
                "path": file_path,
                "hash": file_hash,
                "chunks": len(documents),
                "added_at": datetime.now().isoformat(),
                "chat_id": chat_id
            }
            
            # Save updated vectorstore
            self.save_vectorstore()
            
            return True, f"Successfully processed document: {len(documents)} chunks added to vector database"
        
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            logger.error(traceback.format_exc())
            return False, f"Error processing document: {str(e)}"

    async def extract_text_from_file(self, file_path):
        """Extract text content from various file types"""
        try:
            mime_type, _ = mimetypes.guess_type(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if mime_type and mime_type.startswith('text') or file_extension == '.md':
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return await f.read()
                    
            elif file_extension == '.csv':
                try:
                    df = pd.read_csv(file_path)
                    return df.to_string()
                except Exception as e:
                    logger.error(f"Error reading CSV: {e}")
                    async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return await f.read()
                    
            elif mime_type and mime_type.startswith('application/pdf') or file_extension == '.pdf':
                return self.extract_pdf_text(file_path)
                
            elif mime_type and mime_type.startswith('image'):
                img = Image.open(file_path)
                return pytesseract.image_to_string(img)
                
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                return df.to_string()
                
            elif file_extension == '.docx':
                doc = docx.Document(file_path)
                return '\n\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
            else:
                logger.warning(f"Unsupported file type: {mime_type or file_extension}")
                return f"[Unsupported file type: {file_extension}]"
                
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return None

    def extract_pdf_text(self, file_path):
        """Extract text from PDF with formatting"""
        try:
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num, page in enumerate(doc, 1):
                text_parts.append(f"\n## Page {page_num}\n")
                
                # Get text with block information to preserve some structure
                blocks = page.get_text("blocks")
                for block in blocks:
                    text = block[4].strip()
                    if text:
                        text_parts.append(text + "\n")
                
                text_parts.append("\n---\n")
            
            return ''.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None

    def retrieve(self, query, k=5, filter_chat_id=None):
        """Retrieve relevant document chunks for a query"""
        try:
            # Apply filter if chat_id is provided
            search_filter = None
            if filter_chat_id:
                search_filter = {"chat_id": filter_chat_id}
                
            # Perform the similarity search
            docs_and_scores = self.vectorstore.similarity_search_with_score(query, k=k, filter=search_filter)
            
            # Format and return the results
            results = []
            for doc, score in docs_and_scores:
                results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "score": float(score)  # Convert score to float for JSON serialization
                })
            
            return results
        except Exception as e:
            logger.error(f"Error retrieving documents: {e}")
            return []

    def get_document_list(self, chat_id=None):
        """Get list of documents, optionally filtered by chat_id"""
        try:
            documents = []
            for doc_id, meta in self.document_metadata.items():
                if chat_id is None or meta.get("chat_id") == chat_id:
                    documents.append({
                        "id": doc_id,
                        "filename": meta["filename"],
                        "chunks": meta["chunks"],
                        "added_at": meta["added_at"]
                    })
            return documents
        except Exception as e:
            logger.error(f"Error getting document list: {e}")
            return []

    async def refresh_document_index(self):
        """Refresh document index after deletion"""
        try:
            logger.info("Refreshing document index")
            # Rebuild the document index if needed
            self.save_vectorstore()
            return True
        except Exception as e:
            logger.error(f"Error refreshing document index: {e}")
            logger.error(traceback.format_exc())
            return False

    def delete_document(self, doc_id):
        """Delete a document from the vectorstore"""
        try:
            if doc_id in self.document_metadata:
                # Get all document IDs in the docstore
                docstore_ids = self.vectorstore.docstore._dict.keys()
                docs_to_keep = []
                
                # Get all documents that are not part of the document to delete
                for docstore_id in docstore_ids:
                    doc = self.vectorstore.docstore.search(docstore_id)
                    if doc.metadata.get("doc_id") != doc_id:
                        docs_to_keep.append(doc)
                
                # Create a new vectorstore with only the documents we want to keep
                if docs_to_keep:
                    new_vectorstore = FAISS.from_documents(docs_to_keep, embeddings_model)
                    self.vectorstore = new_vectorstore
                else:
                    # If no documents left, initialize with an empty document
                    self.vectorstore = FAISS.from_documents(
                        [Document(page_content="Initialization document", metadata={"source": "init"})], 
                        embeddings_model
                    )
                
                # Remove from metadata
                if doc_id in self.document_metadata:
                    path = self.document_metadata[doc_id].get("path")
                    del self.document_metadata[doc_id]
                
                # Save the updated vectorstore
                self.save_vectorstore()
                
                return True, f"Document {doc_id} deleted successfully"
            else:
                return False, f"Document {doc_id} not found"
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            logger.error(traceback.format_exc())
            return False, f"Error deleting document: {str(e)}"

    def query_ollama(self, prompt, model=None):
        """Send a query to Ollama API"""
        try:
            url = "http://localhost:11434/api/generate"
            data = {
                "model": model if model else self.model_name,
                "prompt": prompt,
                "stream": False
            }
            response = requests.post(url, json=data)
            if response.status_code == 200:
                return json.loads(response.text)["response"]
            else:
                raise Exception(f"Ollama API error: {response.status_code}, {response.text}")
        except Exception as e:
            logger.error(f"Error in query_ollama: {e}")
            raise

    def generate(self, query, chat_history=None, model=None, use_rag=True, system_prompt=None, include_memory=True):
        """Generate a response to a query, using RAG if enabled and personal memory if available"""
        try:
            # Build context from different sources
            context_parts = []
            
            # 1. Get RAG context if enabled
            if use_rag:
                # Retrieve relevant documents
                retrieved_docs = self.retrieve(query, k=10)
                
                # Format context from retrieved documents
                rag_context_parts = []
                for i, doc in enumerate(retrieved_docs):
                    source = doc["metadata"].get("filename", "unknown")
                    rag_context_parts.append(f"[Document {i+1} from {source}]:\n{doc['content']}\n")
                
                if rag_context_parts:
                    context_parts.append("Context from relevant documents:\n" + "\n\n".join(rag_context_parts))
            
            # 2. Add personal memory context if enabled
            if include_memory and memory_system is not None:
                try:
                    # Get relevant memories for the query
                    memories = asyncio.run(memory_system.get_relevant_memories(query, k=3))
                    
                    if memories:
                        memory_context = memory_system.format_memories_for_context(memories)
                        context_parts.append(memory_context)
                except Exception as mem_err:
                    logger.error(f"Error getting personal memories: {mem_err}")
            
            # Combine all context parts
            context = "\n\n".join(context_parts)
            
            # Format conversation history if provided
            conversation = ""
            if chat_history and len(chat_history) > 0:
                history_parts = []
                for entry in chat_history:
                    role = "Human" if entry.get("is_user") else "Assistant"
                    content = entry.get("content", "")
                    history_parts.append(f"{role}: {content}")
                conversation = "\n".join(history_parts[-10:])  # Use last 10 messages for context
            
            # Construct the prompt
            if conversation:
                if context:
                    prompt = f"""{system_prompt or "You are a helpful assistant that answers based on the provided context."}
                    
Previous conversation:
{conversation}

{context}

Human question: {query}

Using the provided context, answer the human's question. If context contains personal information about the user, use it appropriately to personalize your response. For information from documents, always refer to the document content when answering, even if the match isn't perfect. If no exact information is found, try to extract relevant details that might help answer the question. If you are absolutely certain that the context doesn't contain relevant information, only then say you don't have enough information about this specific topic."""
                else:
                    prompt = f"""{system_prompt or "You are a helpful assistant."}

Previous conversation:
{conversation}

Human question: {query}

Answer the human's question directly and concisely."""
            else:
                if context:
                    prompt = f"""{system_prompt or "You are a helpful assistant that answers based on the provided context."}
                    
{context}

Human question: {query}

Using the provided context, answer the human's question. If context contains personal information about the user, use it appropriately to personalize your response. For information from documents, always refer to the document content when answering, even if the match isn't perfect. If no exact information is found, try to extract relevant details that might help answer the question. If you are absolutely certain that the context doesn't contain relevant information, only then say you don't have enough information about this specific topic."""
                else:
                    prompt = f"""{system_prompt or "You are a helpful assistant."}

Human question: {query}

Answer the human's question directly and concisely."""
            
            # Send the prompt to Ollama
            return self.query_ollama(prompt, model=model)
        except Exception as e:
            logger.error(f"Error generating response: {e}")
            return f"I'm sorry, I encountered an error: {str(e)}"

    def _build_calendar_prompt(self, query, chat_history=None, system_prompt=None):
        """Build a prompt to help the LLM understand calendar intentions"""
        current_date = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        # Start with system instructions
        base_prompt = system_prompt or "You are a helpful assistant with calendar management capabilities."
        
        calendar_instructions = f"""
You are processing a calendar-related request. The current date and time is {current_date}.

TASK: Analyze the user's message to identify the calendar operation they want to perform.
Extract all relevant details like event title, date, time, location, and description.

IMPORTANT: The user may write in different languages. Make sure to correctly understand their intent regardless of language.

Output a JSON object with the following structure:
{{
  "intent": "create"|"update"|"delete"|"query",
  "title": "Event title",
  "date": "YYYY-MM-DD", // Optional, for specific date queries
  "start_time": "YYYY-MM-DDTHH:MM:SS", // ISO format with timezone if available
  "end_time": "YYYY-MM-DDTHH:MM:SS", // ISO format with timezone if available
  "description": "Event description", // Optional
  "location": "Event location", // Optional
  "category": "default"|"work"|"personal"|"holiday"|"other", // Optional
  
  // For update operations only:
  "original_title": "Original event title to update", // Only for update intent
  "new_title": "New event title", // Only for update intent when changing title
  
  // For query operations:
  "start_date": "YYYY-MM-DD", // Optional, for date range queries
  "end_date": "YYYY-MM-DD" // Optional, for date range queries
}}

IMPORTANT RULES:
1. Parse dates intelligently, converting relative dates like "tomorrow", "today", etc. to ISO dates.
2. Set default meeting duration to 1 hour if not specified.
3. For morning events without specific time, default to 9:00 AM.
4. For afternoon events without specific time, default to 2:00 PM.
5. For evening events without specific time, default to 6:00 PM.
6. Time mentions should be properly parsed regardless of format.
7. Include only the fields that are relevant to the specific intent.
8. Return VALID JSON that can be parsed programmatically.
"""

# Import CalendarRAGIntegration for later use
from calendar_integration import CalendarRAGIntegration

async def init_db():
    async with aiosqlite.connect('db/assistant.db') as db:
        await db.execute('''
            CREATE TABLE IF NOT EXISTS chats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at TIMESTAMP,
                title TEXT,
                "order" INTEGER,
                pinned BOOLEAN DEFAULT 0,
                manually_renamed BOOLEAN DEFAULT 0
            )
        ''')
        await db.execute('''
            CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id INTEGER,
                content TEXT,
                is_user BOOLEAN,
                timestamp TIMESTAMP,
                FOREIGN KEY (chat_id) REFERENCES chats (id)
            )
        ''')
        await db.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id INTEGER,
                filename TEXT,
                file_path TEXT,
                doc_id TEXT,
                added_at TIMESTAMP,
                chunks INTEGER,
                FOREIGN KEY (chat_id) REFERENCES chats (id)
            )
        ''')
        
        # Memory system tables
        await db.execute('''
            CREATE TABLE IF NOT EXISTS personal_memory (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                category TEXT NOT NULL,
                content TEXT NOT NULL,
                source_text TEXT NOT NULL,
                confidence REAL NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                last_accessed TIMESTAMP,
                access_count INTEGER DEFAULT 0,
                chat_id INTEGER,
                importance REAL DEFAULT 0.5
            )
        ''')
        
        await db.execute('''
            CREATE TABLE IF NOT EXISTS memory_settings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                setting_name TEXT UNIQUE NOT NULL,
                setting_value TEXT NOT NULL,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Insert default memory settings if they don't exist
        await db.execute('''
            INSERT OR IGNORE INTO memory_settings (setting_name, setting_value)
            VALUES (?, ?)
        ''', ('memory_enabled', 'true'))
        
        await db.execute('''
            INSERT OR IGNORE INTO memory_settings (setting_name, setting_value)
            VALUES (?, ?)
        ''', ('extraction_threshold', '0.7'))
        
        await db.commit()

async def create_chat():
    try:
        async with aiosqlite.connect('db/assistant.db') as db:
            cursor = await db.execute('SELECT COUNT(*) FROM chats')
            count = (await cursor.fetchone())[0]
            await db.execute('INSERT INTO chats (created_at, title, "order") VALUES (?, ?, ?)', 
                            (datetime.now(), 'New Chat', count))
            await db.commit()
            cursor = await db.execute('SELECT last_insert_rowid()')
            return (await cursor.fetchone())[0]
    except sqlite3.OperationalError as e:
        logger.error(f"Database error: {e}")
        await init_db()  # Attempt to initialize the DB
        return await create_chat()  # Retry creating the chat

async def store_message(chat_id, content, is_user):
    try:
        async with aiosqlite.connect('db/assistant.db') as db:
            cursor = await db.execute('''
                INSERT INTO messages (chat_id, content, is_user, timestamp)
                VALUES (?, ?, ?, ?)
            ''', (chat_id, content, is_user, datetime.now()))
            await db.commit()
            message_id = cursor.lastrowid
            
            # Process message for personal information if it's from the user and memory system is initialized
            if is_user and memory_system is not None:
                try:
                    # Extract and store personal information
                    extracted_count = await memory_system.process_message(content, chat_id, is_user)
                    if extracted_count:
                        logger.info(f"Extracted {extracted_count} personal facts from message")
                except Exception as mem_err:
                    logger.error(f"Error processing message for memory: {mem_err}")
            
            return message_id
    except sqlite3.OperationalError as e:
        logger.error(f"Database error in store_message: {e}")
        await init_db()
        return await store_message(chat_id, content, is_user)

async def get_chat_history(chat_id, limit=20):
    try:
        async with aiosqlite.connect('db/assistant.db') as db:
            async with db.execute('''
                SELECT id, content, is_user FROM messages
                WHERE chat_id = ?
                ORDER BY timestamp DESC LIMIT ?
            ''', (chat_id, limit)) as cursor:
                messages = await cursor.fetchall()
                formatted_messages = [{'id': msg[0], 'content': msg[1], 'is_user': bool(msg[2])} for msg in messages]
                return formatted_messages
    except sqlite3.OperationalError as e:
        logger.error(f"Database error in get_chat_history: {e}")
        await init_db()
        return await get_chat_history(chat_id, limit)

async def update_chat_title(chat_id, title, manually_renamed=False):
    try:
        async with aiosqlite.connect('db/assistant.db') as db:
            await db.execute('UPDATE chats SET title = ?, manually_renamed = ? WHERE id = ?', (title, manually_renamed, chat_id))
            await db.commit()
    except sqlite3.OperationalError as e:
        logger.error(f"Database error in update_chat_title: {e}")
        await init_db()
        await update_chat_title(chat_id, title, manually_renamed)

async def generate_chat_title(chat_id):
    try:
        async with aiosqlite.connect('db/assistant.db') as db:
            async with db.execute('SELECT manually_renamed FROM chats WHERE id = ?', (chat_id,)) as cursor:
                result = await cursor.fetchone()
                if result and result[0]:
                    return None

        chat_history = await get_chat_history(chat_id, limit=5)
        prompt = "Based on the following chat messages, generate a brief, descriptive title for this conversation. Do not use any Markdown formatting in the title - return plain text only, no asterisks, no bold or other formatting:\n\n"
        for msg in reversed(chat_history):
            role = "Human" if msg['is_user'] else "Assistant"
            prompt += f"{role}: {msg['content']}\n"
        prompt += "\nTitle:"

        title = rag_system.query_ollama(prompt)
        await update_chat_title(chat_id, title)
        return title
    except sqlite3.OperationalError as e:
        logger.error(f"Database error in generate_chat_title: {e}")
        await init_db()
        return await generate_chat_title(chat_id)

# Helper functions
def get_chat_upload_folder(chat_id):
    chat_folder = os.path.join(UPLOAD_FOLDER, str(chat_id))
    os.makedirs(chat_folder, exist_ok=True)
    return chat_folder

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

async def save_uploaded_file(file, chat_id):
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        chat_folder = get_chat_upload_folder(chat_id)
        os.makedirs(chat_folder, exist_ok=True)
        file_path = os.path.join(chat_folder, filename)
        await file.save(file_path)
        return file_path
    return None

async def perform_duckduckgo_search(query, max_results=5):
    results = ddg(query, max_results=max_results)
    return [{'title': result.get('title', ''), 'link': result.get('href', ''), 'snippet': result.get('body', '')} for result in results]

def generate_duckduckgo_summary(query, results):
    summary = f"I performed a DuckDuckGo search for '{query}'. Here's a summary of the top results:\n\n"
    for i, result in enumerate(results, 1):
        summary += f"{i}. {result['title']}\n"
        summary += f"   URL: {result['link']}\n"
        summary += f"   Description: {result['snippet']}\n\n"
    return summary

async def process_image_with_llava(image_path):
    with open(image_path, "rb") as image_file:
        image_data = base64.b64encode(image_file.read()).decode('utf-8')
        
    prompt = "Analyze this image and describe what you see in detail."
    
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post('http://localhost:11434/api/generate', 
                                    json={
                                        "model": image_recognition_model,
                                        "prompt": prompt,
                                        "images": [image_data],
                                        "stream": False
                                    }) as response:
                response.raise_for_status()
                response_data = await response.json()
            
        if 'response' in response_data:
            return response_data['response']
        else:
            return "I'm sorry, I couldn't generate a description for this image."
    except Exception as e:
        logger.error(f"Error in image recognition: {e}")
        return "I'm sorry, I encountered an error while processing the image."

# API Routes
@app.route('/available_models', methods=['GET'])
async def get_available_models():
    try:
        async with aiohttp.ClientSession() as session:
            # Using the Ollama API endpoint to get all available models
            async with session.get('http://localhost:11434/api/tags') as response:
                response.raise_for_status()
                response_data = await response.json()
                
        if 'models' in response_data:
            models = [model['name'] for model in response_data['models']]
            return jsonify({'models': models})
        else:
            logger.error("No models field in Ollama API response")
            return jsonify({'models': [llm_model, image_recognition_model]})
    except Exception as e:
        logger.error(f"Error fetching models: {e}")
        return jsonify({'models': [llm_model, image_recognition_model]})

@app.route('/rag/documents', methods=['GET'])
async def get_documents():
    try:
        chat_id = request.args.get('chat_id')
        chat_id = int(chat_id) if chat_id else None
        
        documents = rag_system.get_document_list(chat_id)
        return jsonify({'documents': documents})
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        return jsonify({'error': 'An error occurred while fetching documents'}), 500

@app.route('/rag/documents/<doc_id>', methods=['DELETE'])
async def delete_document(doc_id):
    try:
        # Get document metadata before deletion to find the file path
        doc_metadata = rag_system.document_metadata.get(doc_id, {})
        document_path = doc_metadata.get('path')
        
        # Removing vector database files
        faiss_path = os.path.join(VECTOR_DB_FOLDER, f"{doc_id}.faiss")
        index_path = os.path.join(VECTOR_DB_FOLDER, f"{doc_id}.pkl")
        
        files_deleted = 0
        if os.path.exists(faiss_path):
            os.remove(faiss_path)
            files_deleted += 1
            logger.info(f"Deleted vector file: {faiss_path}")
        
        if os.path.exists(index_path):
            os.remove(index_path)
            files_deleted += 1
            logger.info(f"Deleted index file: {index_path}")
        
        # Remove from RAG system
        rag_system.delete_document(doc_id)
        
        # Delete from database
        async with aiosqlite.connect('db/assistant.db') as db:
            # Check if the documents table exists
            async with db.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='documents'") as cursor:
                if await cursor.fetchone():
                    await db.execute('DELETE FROM documents WHERE doc_id = ?', (doc_id,))
                    logger.info(f"Deleted document {doc_id} from documents table")
            
            # Also check for the rag_documents table
            async with db.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='rag_documents'") as cursor:
                if await cursor.fetchone():
                    await db.execute('DELETE FROM rag_documents WHERE id = ?', (doc_id,))
                    logger.info(f"Deleted document {doc_id} from rag_documents table")
            
            await db.commit()
        
        # Delete the actual document file if it exists and is in uploads directory
        if document_path and os.path.exists(document_path) and '/uploads/' in document_path:
            try:
                os.remove(document_path)
                files_deleted += 1
                logger.info(f"Deleted original file: {document_path}")
            except Exception as file_error:
                logger.warning(f"Could not delete original file: {file_error}")
        
        logger.info(f"Document {doc_id} deleted successfully")
        return jsonify({'success': True, 'message': f'Document deleted successfully. {files_deleted} files removed'})
    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': 'An error occurred while deleting the document'}), 500

@app.route('/rag/search', methods=['POST'])
async def search_documents():
    try:
        data = await request.get_json()
        query = data.get('query')
        chat_id = data.get('chat_id')
        limit = data.get('limit', 5)
        
        if not query:
            return jsonify({'error': 'No query provided'}), 400
            
        chat_id = int(chat_id) if chat_id else None
        results = rag_system.retrieve(query, k=limit, filter_chat_id=chat_id)
        
        return jsonify({'results': results})
    except Exception as e:
        logger.error(f"Error searching documents: {str(e)}")
        return jsonify({'error': 'An error occurred while searching documents'}), 500

@app.route('/rag/upload', methods=['POST'])
async def upload_document():
    try:
        form = await request.form
        files = await request.files
        if 'file' not in files:
            return jsonify({"error": "No file part"}), 400
        
        file = files['file']
        chat_id = form.get('chat_id')
        client_system_prompt = form.get('system_prompt')
        current_system_prompt = client_system_prompt if client_system_prompt else system_prompt
        
        if not chat_id:
            chat_id = await create_chat()
        else:
            chat_id = int(chat_id)
            
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        
        file_path = await save_uploaded_file(file, chat_id)
        
        if file_path:
            # Process the document with RAG
            success, message = await rag_system.process_document(file_path, chat_id)
            
            if success:
                # Store document in database
                doc_id = f"{os.path.basename(file_path)}_{rag_system.get_file_hash(file_path)}"
                async with aiosqlite.connect('db/assistant.db') as db:
                    await db.execute('''
                        INSERT INTO documents (chat_id, filename, file_path, doc_id, added_at, chunks)
                        VALUES (?, ?, ?, ?, ?, ?)
                    ''', (
                        chat_id, 
                        os.path.basename(file_path), 
                        file_path, 
                        doc_id, 
                        datetime.now(),
                        rag_system.document_metadata.get(doc_id, {}).get('chunks', 0)
                    ))
                    await db.commit()
                    
                return jsonify({
                    "success": True,
                    "message": message,
                    "file_path": file_path,
                    "chat_id": chat_id,
                    "doc_id": doc_id,
                    "content": await rag_system.extract_text_from_file(file_path)
                })
            else:
                return jsonify({"error": message}), 400
        
        return jsonify({"error": "File type not allowed"}), 400
    
    except Exception as e:
        logger.error(f"Error processing document upload: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500

@app.route('/chat', methods=['POST'])
async def chat():
    global system_prompt
    try:
        data = await request.get_json()
        message = data.get('message')
        chat_id = data.get('chat_id')
        use_duckduckgo = data.get('use_duckduckgo', False)
        use_rag = data.get('use_rag', True)
        selected_model = data.get('selected_model', llm_model)
        client_system_prompt = data.get('system_prompt')
        use_deep_search = data.get('use_deep_search', False)
        use_memory = data.get('use_memory', True)
        
        # Use system prompt from client if it's provided
        current_system_prompt = client_system_prompt if client_system_prompt else system_prompt
        
        if not message:
            return jsonify({'error': 'No message provided'}), 400
        
        if not chat_id:
            chat_id = await create_chat()
        else:
            chat_id = int(chat_id)
            
        # Get chat history for context
        chat_history = await get_chat_history(chat_id, limit=10)
        chat_history.reverse()  # Most recent last
        
        # Check for memory management commands
        if message.startswith("/memory"):
            if memory_system is None:
                memory_response = "Memory system is not initialized yet. Please try again in a moment."
            else:
                memory_response = await memory_system.process_memory_command(message)
            
            # Store command and response in chat history
            user_message_id = await store_message(chat_id, message, True)
            assistant_message_id = await store_message(chat_id, memory_response, False)
            
            # Generate chat title if needed
            chat_title = await generate_chat_title(chat_id)
            
            return jsonify({
                'response': memory_response,
                'chat_id': chat_id,
                'chat_title': chat_title,
                'user_message_id': user_message_id,
                'assistant_message_id': assistant_message_id,
                'update_calendar': False,
                'is_memory_command': True
            })
        
        # Handle Deep Search case separately
        if use_deep_search:
            logger.info(f"Handling chat request with Deep Search for query: {message}")
            # Run the entire deep search process from search.py
            response = perform_deep_search(message)
            
            # Store messages
            user_message_id = await store_message(chat_id, message, True)
            assistant_message_id = await store_message(chat_id, response, False)
            chat_title = await generate_chat_title(chat_id)
            
            return jsonify({
                'response': response,
                'chat_id': chat_id,
                'chat_title': chat_title,
                'user_message_id': user_message_id,
                'assistant_message_id': assistant_message_id,
                'update_calendar': False
            })

        # --- Handle non-Deep Search cases (Calendar, DuckDuckGo, RAG) ---
        
        # Check if this is a calendar-related query using a basic keyword check
        calendar_keywords = [
            'schedule', 'appointment', 'meeting', 'event', 'calendar',
            'add to calendar', 'create event', 'remove event', 'delete event',
            'cancel meeting', 'reschedule', 'move meeting',
            'show calendar', 'list events', 'upcoming events',
            'tomorrow', 'next week', 'sync calendar', 'reminder', 'remind me',
            'plan', 'agenda', 'booking', 'reservation',
            # Russian keywords
            'календарь', 'встреча', 'встречи', 'событие', 'события',
            'запланировано', 'расписание', 'напомни', 'напоминание',
            'запись', 'записи', 'создай', 'добавь', 'удали', 'убери',
            'отмени', 'перенеси', 'покажи', 'какие', 'когда', 'завтра',
            'сегодня', 'план', 'планы'
        ]
        
        is_calendar_query = any(keyword in message.lower() for keyword in calendar_keywords)
        
        if is_calendar_query:
            logger.info(f"Detected calendar query in chat: {message}")
            # Process calendar query
            calendar_result = await calendar_rag.process_calendar_query(
                query=message,
                chat_id=chat_id,
                selected_model=selected_model,
                system_prompt=current_system_prompt
            )
            
            logger.info(f"Calendar result in chat: {json.dumps(calendar_result)}")
            
            # Store messages in database
            user_message_id = await store_message(chat_id, message, True)
            assistant_message_id = await store_message(chat_id, calendar_result.get('message', "I processed your calendar request."), False)
            
            # Generate chat title if needed
            chat_title = await generate_chat_title(chat_id)
            
            return jsonify({
                'response': calendar_result.get('message', "I processed your calendar request."),
                'chat_id': chat_id,
                'chat_title': chat_title,
                'user_message_id': user_message_id,
                'assistant_message_id': assistant_message_id,
                'update_calendar': True,
                'calendar_action': calendar_result.get('intent'),
                'success': calendar_result.get('success', False)
            })
    
        # Prepare context for the LLM (only for DuckDuckGo or RAG)
        final_system_prompt = current_system_prompt
    
        # Add DuckDuckGo search results if requested
        if use_duckduckgo:
            search_results = await perform_duckduckgo_search(message)
            search_summary = generate_duckduckgo_summary(message, search_results)
            final_system_prompt += f"\n\nContext from DuckDuckGo search: {search_summary}\n"
            
        # Generate response using RAG (Deep Search case is handled above)
        response = rag_system.generate(
            query=message,
            chat_history=chat_history,
            model=selected_model,
            use_rag=use_rag,
            system_prompt=final_system_prompt, # Use the potentially enhanced prompt
            include_memory=use_memory  # Use memory if enabled
        )
    
        # Store messages in database
        user_message_id = await store_message(chat_id, message, True)
        assistant_message_id = await store_message(chat_id, response, False)
    
        # Generate chat title if needed
        chat_title = await generate_chat_title(chat_id)
    
        # Include search results in response if DuckDuckGo was used
        if use_duckduckgo:
            response += "\n\nHere are the original DuckDuckGo search results I based my summary on:\n\n"
            for item in search_results:
                response += f"- {item['title']}\n  {item['link']}\n"
                
        return jsonify({
            'response': response, 
            'chat_id': chat_id, 
            'chat_title': chat_title,
            'user_message_id': user_message_id,
            'assistant_message_id': assistant_message_id,
            'update_calendar': False
        })
        
    except Exception as e:
        logger.error(f"Error processing chat request: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500

@app.route('/upload', methods=['POST'])
async def upload_file():
    try:
        form = await request.form
        files = await request.files
        if 'file' not in files:
            return jsonify({"error": "No file part"}), 400
        
        file = files['file']
        chat_id = form.get('chat_id')
        selected_model = form.get('selected_model', llm_model)
        client_system_prompt = form.get('system_prompt')
        current_system_prompt = client_system_prompt if client_system_prompt else system_prompt
        
        if not chat_id:
            chat_id = await create_chat()
        else:
            chat_id = int(chat_id)
            
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        
        file_path = await save_uploaded_file(file, chat_id)
        
        if file_path:
            mime_type, _ = mimetypes.guess_type(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Process file for RAG indexing in background
            success, process_message = await rag_system.process_document(file_path, chat_id)
            
            # Get file content for message
            if mime_type and mime_type.startswith('image'):
                file_content = await process_image_with_llava(file_path)
            else:
                file_content = await rag_system.extract_text_from_file(file_path)
                # Truncate if too long
                if file_content and len(file_content) > 2000:
                    file_content = file_content[:2000] + "...\n[Content truncated for display]"
                    
            # Get chat history for context
            chat_history = await get_chat_history(chat_id, limit=5)
            chat_history.reverse()
            
            # Generate response with RAG
            response = rag_system.generate(
                query=f"The user has uploaded a file named {file.filename}. The content is: {file_content[:1000]}... Provide a helpful analysis of this content.",
                chat_history=chat_history,
                model=selected_model,
                use_rag=False,  # Don't use RAG for the initial file upload response
                system_prompt=current_system_prompt
            )
            
            # Add information about the RAG processing
            if success:
                response += f"\n\n(The document has been successfully indexed with {rag_system.document_metadata.get(f'{os.path.basename(file_path)}_{rag_system.get_file_hash(file_path)}', {}).get('chunks', 0)} chunks. You can now ask questions about this document.)"
                
            # Store messages in database
            user_message_id = await store_message(chat_id, f"I've uploaded a file named {file.filename}. Please analyze this document and help me understand its content.", True)
            assistant_message_id = await store_message(chat_id, response, False)
            
            # Generate chat title if needed
            chat_title = await generate_chat_title(chat_id)
            
            return jsonify({
                "response": response, 
                "chat_id": chat_id, 
                "chat_title": chat_title,
                "user_message_id": user_message_id,
                "assistant_message_id": assistant_message_id,
                "file_path": f"/uploads/{chat_id}/{os.path.basename(file_path)}",
                "mime_type": mime_type,
                "file_content": file_content,
                "rag_status": "Indexed successfully" if success else "Indexing failed"
            })
        
        return jsonify({"error": "File type not allowed"}), 400
    
    except Exception as e:
        logger.error(f"Error processing file upload: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': 'An error occurred while processing your file upload. Please try again.'}), 500
    
@app.route('/uploads/<int:chat_id>/<path:filename>')
async def uploaded_file(chat_id, filename):
    return await send_file(os.path.join(get_chat_upload_folder(chat_id), filename))

@app.route('/chat_history', methods=['GET'])
async def get_chats():
    try:
        async with aiosqlite.connect('db/assistant.db') as db:
            async with db.execute('SELECT id, created_at, title, "order", pinned, manually_renamed FROM chats ORDER BY pinned DESC, "order" ASC, created_at DESC') as cursor:
                chats = [{'id': row[0], 'created_at': row[1], 'title': row[2], 'order': row[3], 'pinned': bool(row[4]), 'manually_renamed': bool(row[5])} for row in await cursor.fetchall()]
        return jsonify({'chats': chats})
    except Exception as e:
        logger.error(f"Error fetching chat history: {str(e)}")
        return jsonify({'error': 'An error occurred while fetching chat history'}), 500
    
@app.route('/chat/<int:chat_id>', methods=['GET'])
async def get_chat(chat_id):
    try:
        messages = await get_chat_history(chat_id, limit=100)
        # Get documents related to this chat
        documents = rag_system.get_document_list(chat_id)
        return jsonify({
            'messages': messages,
            'documents': documents
        })
    except Exception as e:
        logger.error(f"Error fetching chat: {str(e)}")
        return jsonify({'error': 'An error occurred while fetching the chat'}), 500
    
@app.route('/reorder_chats', methods=['POST'])
async def reorder_chats():
    try:
        data = await request.get_json()
        new_order = data['new_order']
        
        async with aiosqlite.connect('db/assistant.db') as db:
            for index, chat_id in enumerate(new_order):
                await db.execute('UPDATE chats SET "order" = ? WHERE id = ?', (index, chat_id))
            await db.commit()
            
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error reordering chats: {str(e)}")
        return jsonify({'error': 'An error occurred while reordering chats'}), 500
    
@app.route('/toggle_pin', methods=['POST'])
async def toggle_pin():
    try:
        data = await request.get_json()
        chat_id = data['chat_id']
        
        async with aiosqlite.connect('db/assistant.db') as db:
            await db.execute('UPDATE chats SET pinned = NOT pinned WHERE id = ?', (chat_id,))
            await db.commit()
            
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error toggling pin: {str(e)}")
        return jsonify({'error': 'An error occurred while toggling pin status'}), 500
    
@app.route('/delete_chat/<int:chat_id>', methods=['DELETE'])
async def delete_chat(chat_id):
    try:
        # First, get all documents in this chat to remove them from RAG
        async with aiosqlite.connect('db/assistant.db') as db:
            async with db.execute('SELECT doc_id FROM documents WHERE chat_id = ?', (chat_id,)) as cursor:
                documents = await cursor.fetchall()
                
            # Delete each document from the RAG system
            for doc in documents:
                doc_id = doc[0]
                rag_system.delete_document(doc_id)
                
            # Delete from database
            await db.execute('DELETE FROM documents WHERE chat_id = ?', (chat_id,))
            await db.execute('DELETE FROM messages WHERE chat_id = ?', (chat_id,))
            await db.execute('DELETE FROM chats WHERE id = ?', (chat_id,))
            await db.commit()
            
        # Delete uploaded files
        chat_folder = get_chat_upload_folder(chat_id)
        if os.path.exists(chat_folder):
            shutil.rmtree(chat_folder)
            
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error deleting chat: {str(e)}")
        return jsonify({'error': 'An error occurred while deleting the chat'}), 500
    
@app.route('/delete_message', methods=['POST'])
async def delete_message():
    try:
        data = await request.get_json()
        message_id = data['message_id']
        
        async with aiosqlite.connect('db/assistant.db') as db:
            await db.execute('DELETE FROM messages WHERE id = ?', (message_id,))
            await db.commit()
            
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error deleting message: {str(e)}")
        return jsonify({'error': 'An error occurred while deleting the message'}), 500
    
@app.route('/update_settings', methods=['POST'])
async def update_settings():
    """Update application settings and Google Calendar settings"""
    try:
        data = await request.get_json()
        
        # Update LLM and image models
        global llm_model, image_recognition_model, system_prompt
        if 'llm_model' in data:
            llm_model = data['llm_model']
        if 'image_model' in data:
            image_recognition_model = data['image_model']
        if 'system_prompt' in data:
            system_prompt = data['system_prompt']
        
        # Update Google Calendar settings in .env file
        if 'google_api_key' in data or 'google_client_id' in data or 'google_client_secret' in data or 'gmail_api_key' in data or 'gmail_client_id' in data or 'gmail_client_secret' in data:
            # Read current .env file
            env_content = []
            if os.path.exists('.env'):
                with open('.env', 'r') as f:
                    env_content = f.readlines()
            
            # Find and update or add settings
            updated_google_api_key = False
            updated_google_client_id = False
            updated_google_client_secret = False
            updated_gmail_api_key = False
            updated_gmail_client_id = False
            updated_gmail_client_secret = False
            
            for i, line in enumerate(env_content):
                if line.startswith('GOOGLE_API_KEY='):
                    if 'google_api_key' in data:
                        env_content[i] = f"GOOGLE_API_KEY={data['google_api_key']}\n"
                        updated_google_api_key = True
                elif line.startswith('GOOGLE_CLIENT_ID='):
                    if 'google_client_id' in data:
                        env_content[i] = f"GOOGLE_CLIENT_ID={data['google_client_id']}\n"
                        updated_google_client_id = True
                elif line.startswith('GOOGLE_CLIENT_SECRET='):
                    if 'google_client_secret' in data:
                        env_content[i] = f"GOOGLE_CLIENT_SECRET={data['google_client_secret']}\n"
                        updated_google_client_secret = True
                elif line.startswith('GMAIL_API_KEY='):
                    if 'gmail_api_key' in data:
                        env_content[i] = f"GMAIL_API_KEY={data['gmail_api_key']}\n"
                        updated_gmail_api_key = True
                elif line.startswith('GMAIL_CLIENT_ID='):
                    if 'gmail_client_id' in data:
                        env_content[i] = f"GMAIL_CLIENT_ID={data['gmail_client_id']}\n"
                        updated_gmail_client_id = True
                elif line.startswith('GMAIL_CLIENT_SECRET='):
                    if 'gmail_client_secret' in data:
                        env_content[i] = f"GMAIL_CLIENT_SECRET={data['gmail_client_secret']}\n"
                        updated_gmail_client_secret = True
            
            # Add settings if they don't exist
            google_section_found = False
            gmail_section_found = False
            
            for line in env_content:
                if '# Google Calendar API' in line:
                    google_section_found = True
                if '# Gmail API' in line:
                    gmail_section_found = True
            
            if not google_section_found:
                env_content.append('\n# Google Calendar API\n')
            
            if 'google_api_key' in data and not updated_google_api_key:
                env_content.append(f"GOOGLE_API_KEY={data['google_api_key']}\n")
            if 'google_client_id' in data and not updated_google_client_id:
                env_content.append(f"GOOGLE_CLIENT_ID={data['google_client_id']}\n")
            if 'google_client_secret' in data and not updated_google_client_secret:
                env_content.append(f"GOOGLE_CLIENT_SECRET={data['google_client_secret']}\n")
            
            if not gmail_section_found:
                env_content.append('\n# Gmail API\n')
            
            if 'gmail_api_key' in data and not updated_gmail_api_key:
                env_content.append(f"GMAIL_API_KEY={data['gmail_api_key']}\n")
            if 'gmail_client_id' in data and not updated_gmail_client_id:
                env_content.append(f"GMAIL_CLIENT_ID={data['gmail_client_id']}\n")
            if 'gmail_client_secret' in data and not updated_gmail_client_secret:
                env_content.append(f"GMAIL_CLIENT_SECRET={data['gmail_client_secret']}\n")
            
            # Write updated .env file
            with open('.env', 'w') as f:
                f.writelines(env_content)
            
            # Reload environment variables
            load_dotenv()
        
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error updating settings: {e}")
        return jsonify({'success': False, 'error': str(e)})
    
@app.route('/rename_chat', methods=['POST'])
async def rename_chat():
    try:
        data = await request.get_json()
        chat_id = data.get('chat_id')
        new_title = data.get('new_title')
        if not chat_id or not new_title:
            return jsonify({'error': 'Chat ID and new title are required'}), 400
        
        await update_chat_title(chat_id, new_title, manually_renamed=True)
        return jsonify({'success': True, 'message': 'Chat renamed successfully'})
    except Exception as e:
        logger.error(f"Error renaming chat: {str(e)}")
        return jsonify({'error': 'An error occurred while renaming the chat'}), 500
    
# Frontend-specific routes to add RAG UI components
@app.route('/rag/ui/document_list', methods=['GET'])
async def get_document_list_html():
    """Returns HTML for document list component"""
    try:
        chat_id = request.args.get('chat_id')
        chat_id = int(chat_id) if chat_id else None
        
        documents = rag_system.get_document_list(chat_id)
        
        # Generate simple HTML for document list
        html = """
        <div class="rag-documents">
            <h3>Indexed Documents</h3>
            <ul class="document-list">
        """
        
        if not documents:
            html += "<li>No documents indexed. Upload files to enable RAG.</li>"
        else:
            for doc in documents:
                html += f"""
                <li class="document-item" data-id="{doc['id']}">
                    <span class="doc-name">{doc['filename']}</span>
                    <span class="doc-chunks">({doc['chunks']} chunks)</span>
                    <button class="delete-doc-btn" onclick="deleteDocument('{doc['id']}')">Delete</button>
                </li>
                """
                
        html += """
            </ul>
        </div>
        """
        
        return html
    except Exception as e:
        logger.error(f"Error generating document list: {str(e)}")
        return "<div>Error loading documents</div>"
    
@app.route('/get_settings', methods=['GET'])
async def get_settings():
    try:
        return jsonify({
            'success': True,
            'system_prompt': system_prompt,
            'llm_model': llm_model,
            'image_model': image_recognition_model
        })
    except Exception as e:
        logger.error(f"Error getting settings: {str(e)}")
        return jsonify({'error': 'An error occurred while getting settings'}), 500

# Calendar API Routes
@app.route('/calendar/events', methods=['GET'])
async def get_calendar_events():
    try:
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        search = request.args.get('search')
        
        events = await calendar_manager.get_events(start_date, end_date, search)
        return jsonify({'events': events})
    except Exception as e:
        logger.error(f"Error getting calendar events: {str(e)}")
        return jsonify({'error': 'An error occurred while fetching calendar events'}), 500

@app.route('/calendar/events', methods=['POST'])
async def create_calendar_event():
    try:
        data = await request.get_json()
        required_fields = ['title', 'start_time', 'end_time']
        
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400
                
        # Convert string dates to ISO format if they're not already
        if isinstance(data['start_time'], str):
            try:
                # Ensure it's in ISO format
                datetime.fromisoformat(data['start_time'].replace('Z', '+00:00'))
            except ValueError:
                return jsonify({'error': 'Invalid start_time format, use ISO format'}), 400
                
        if isinstance(data['end_time'], str):
            try:
                # Ensure it's in ISO format
                datetime.fromisoformat(data['end_time'].replace('Z', '+00:00'))
            except ValueError:
                return jsonify({'error': 'Invalid end_time format, use ISO format'}), 400
        
        event_id = await calendar_manager.create_event(data)
        
        if event_id:
            event = await calendar_manager.get_event(event_id)
            return jsonify({'success': True, 'event': event})
        else:
            return jsonify({'error': 'Failed to create event'}), 500
    except Exception as e:
        logger.error(f"Error creating calendar event: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': 'An error occurred while creating the event'}), 500

@app.route('/calendar/events/<int:event_id>', methods=['GET'])
async def get_calendar_event(event_id):
    try:
        event = await calendar_manager.get_event(event_id)
        if event:
            return jsonify({'event': event})
        else:
            return jsonify({'error': 'Event not found'}), 404
    except Exception as e:
        logger.error(f"Error getting calendar event: {str(e)}")
        return jsonify({'error': 'An error occurred while fetching the event'}), 500

@app.route('/calendar/events/<int:event_id>', methods=['PUT'])
async def update_calendar_event(event_id):
    try:
        data = await request.get_json()
        
        # Ensure the event exists
        event = await calendar_manager.get_event(event_id)
        if not event:
            return jsonify({'error': 'Event not found'}), 404
        
        required_fields = ['title', 'start_time', 'end_time']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400
        
        success = await calendar_manager.update_event(event_id, data)
        
        if success:
            updated_event = await calendar_manager.get_event(event_id)
            return jsonify({'success': True, 'event': updated_event})
        else:
            return jsonify({'error': 'Failed to update event'}), 500
    except Exception as e:
        logger.error(f"Error updating calendar event: {str(e)}")
        return jsonify({'error': 'An error occurred while updating the event'}), 500

@app.route('/calendar/events/<int:event_id>', methods=['DELETE'])
async def delete_calendar_event(event_id):
    try:
        # Ensure the event exists
        event = await calendar_manager.get_event(event_id)
        if not event:
            return jsonify({'error': 'Event not found'}), 404
        
        success = await calendar_manager.delete_event(event_id)
        
        if success:
            return jsonify({'success': True, 'message': 'Event deleted successfully'})
        else:
            return jsonify({'error': 'Failed to delete event'}), 500
    except Exception as e:
        logger.error(f"Error deleting calendar event: {str(e)}")
        return jsonify({'error': 'An error occurred while deleting the event'}), 500

@app.route('/calendar/google/auth', methods=['POST'])
async def set_google_calendar_auth():
    try:
        data = await request.get_json()
        
        required_fields = ['access_token', 'refresh_token', 'expiry_time']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400
        
        success = await calendar_manager.save_google_auth(data)
        
        if success:
            return jsonify({'success': True, 'message': 'Google Calendar authentication saved'})
        else:
            return jsonify({'error': 'Failed to save Google Calendar authentication'}), 500
    except Exception as e:
        logger.error(f"Error saving Google Calendar auth: {str(e)}")
        return jsonify({'error': 'An error occurred while saving Google Calendar authentication'}), 500

@app.route('/calendar/google/auth', methods=['GET'])
async def get_google_calendar_auth():
    try:
        auth_data = await calendar_manager.get_google_auth()
        
        if auth_data:
            # Don't return actual tokens for security reasons, just status
            return jsonify({
                'authenticated': True,
                'expiry_time': auth_data['expiry_time']
            })
        else:
            return jsonify({'authenticated': False})
    except Exception as e:
        logger.error(f"Error getting Google Calendar auth: {str(e)}")
        return jsonify({'error': 'An error occurred while getting Google Calendar authentication'}), 500

@app.route('/calendar/natural', methods=['POST'])
async def process_natural_language_calendar():
    try:
        data = await request.get_json()
        message = data.get('message')
        chat_id = data.get('chat_id')
        selected_model = data.get('selected_model')
        
        logger.info(f"Calendar natural language request: {message}")
        
        if not message:
            return jsonify({'error': 'No message provided'}), 400
        
        # Use the CalendarRAGIntegration to process the query
        result = await calendar_rag.process_calendar_query(
            query=message,
            chat_id=chat_id,
            selected_model=selected_model,
            system_prompt=system_prompt
        )
        
        logger.info(f"Calendar processing result: {json.dumps(result)}")
        
        # Store in chat history if chat_id is provided and the request was successful
        user_message_id = None
        assistant_message_id = None
        
        if chat_id:
            chat_id = int(chat_id)
            user_message_id = await store_message(chat_id, message, True)
            assistant_message_id = await store_message(chat_id, result.get('message', ''), False)
        
        # Add message IDs to the result
        result['user_message_id'] = user_message_id
        result['assistant_message_id'] = assistant_message_id
        
        return jsonify(result)
            
    except Exception as e:
        logger.error(f"Error processing natural language calendar command: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': f'An error occurred while processing the calendar command: {str(e)}',
            'message': "I encountered an error while processing your calendar request."
        }), 500

@app.route('/google_calendar_settings', methods=['GET'])
async def get_google_calendar_settings():
    """Get Google Calendar API settings from .env file"""
    try:
        api_key = os.getenv('GOOGLE_API_KEY', '')
        client_id = os.getenv('GOOGLE_CLIENT_ID', '')
        client_secret = os.getenv('GOOGLE_CLIENT_SECRET', '')
        
        return jsonify({
            'success': True,
            'api_key': api_key,
            'client_id': client_id,
            'client_secret': client_secret
        })
    except Exception as e:
        logger.error(f"Error getting Google Calendar settings: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/gmail_settings', methods=['GET'])
async def get_gmail_settings():
    """Get Gmail API settings from .env file"""
    try:
        api_key = os.getenv('GMAIL_API_KEY', '')
        client_id = os.getenv('GMAIL_CLIENT_ID', '')
        client_secret = os.getenv('GMAIL_CLIENT_SECRET', '')
        
        return jsonify({
            'success': True,
            'api_key': api_key,
            'client_id': client_id,
            'client_secret': client_secret
        })
    except Exception as e:
        logger.error(f"Error getting Gmail settings: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/gmail/auth/url', methods=['GET'])
async def get_gmail_auth_url():
    """Get Gmail authorization URL"""
    try:
        # Get Gmail API credentials from .env file
        client_id = os.getenv('GMAIL_CLIENT_ID', '')
        client_secret = os.getenv('GMAIL_CLIENT_SECRET', '')
        
        if not client_id or not client_secret:
            return jsonify({
                'success': False,
                'error': 'Gmail API credentials not configured'
            })
        
        # Create OAuth flow
        oauth_flow = {
            'client_id': client_id,
            'redirect_uri': 'http://localhost:5001/gmail/auth/callback',
            'scope': 'https://www.googleapis.com/auth/gmail.readonly https://www.googleapis.com/auth/gmail.send',
            'response_type': 'code',
            'access_type': 'offline',
            'prompt': 'consent'
        }
        
        # Construct auth URL
        auth_url = f"https://accounts.google.com/o/oauth2/auth?{urllib.parse.urlencode(oauth_flow)}"
        
        return jsonify({
            'success': True,
            'auth_url': auth_url
        })
    except Exception as e:
        logger.error(f"Error generating Gmail auth URL: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/gmail/auth/callback', methods=['GET'])
async def gmail_auth_callback():
    """Handle Gmail auth callback"""
    try:
        # Get auth code from query parameters
        code = request.args.get('code')
        
        if not code:
            return "Authorization failed: No code received", 400
        
        # Get Gmail API credentials from .env file
        client_id = os.getenv('GMAIL_CLIENT_ID', '')
        client_secret = os.getenv('GMAIL_CLIENT_SECRET', '')
        
        if not client_id or not client_secret:
            return "Authorization failed: API credentials not configured", 400
        
        # Exchange code for access token
        token_data = {
            'code': code,
            'client_id': client_id,
            'client_secret': client_secret,
            'redirect_uri': 'http://localhost:5001/gmail/auth/callback',
            'grant_type': 'authorization_code'
        }
        
        token_response = requests.post(
            'https://oauth2.googleapis.com/token',
            data=token_data
        )
        
        if token_response.status_code != 200:
            return f"Authorization failed: {token_response.text}", 400
        
        token_json = token_response.json()
        
        # Store tokens in database
        async with aiosqlite.connect('db/assistant.db') as db:
            await db.execute('''
                CREATE TABLE IF NOT EXISTS gmail_auth_tokens (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    access_token TEXT,
                    refresh_token TEXT,
                    expiry_time INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            expiry_time = int(time.time()) + token_json.get('expires_in', 3600)
            
            # Clear existing tokens
            await db.execute('DELETE FROM gmail_auth_tokens')
            
            # Insert new tokens
            await db.execute('''
                INSERT INTO gmail_auth_tokens 
                (access_token, refresh_token, expiry_time)
                VALUES (?, ?, ?)
            ''', (
                token_json.get('access_token'),
                token_json.get('refresh_token'),
                expiry_time
            ))
            
            await db.commit()
        
        # Return success page with JavaScript to close window and message parent
        return """
        <html>
        <head>
            <title>Gmail Authorization Successful</title>
            <style>
                body {
                    font-family: Arial, sans-serif;
                    text-align: center;
                    padding: 50px;
                    background-color: #f5f5f5;
                }
                .success {
                    color: #4CAF50;
                    font-size: 24px;
                    margin-bottom: 20px;
                }
                .message {
                    margin-bottom: 30px;
                }
                .close-btn {
                    padding: 10px 20px;
                    background-color: #4CAF50;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    cursor: pointer;
                    font-size: 16px;
                }
            </style>
        </head>
        <body>
            <div class="success">Authorization Successful!</div>
            <div class="message">You have successfully authorized VerbaBot to access your Gmail account.</div>
            <button class="close-btn" onclick="closeWindow()">Close Window</button>
            
            <script>
                function closeWindow() {
                    window.opener.postMessage('gmail-auth-success', '*');
                    window.close();
                }
                
                // Auto close after a delay
                setTimeout(function() {
                    window.opener.postMessage('gmail-auth-success', '*');
                    window.close();
                }, 3000);
            </script>
        </body>
        </html>
        """
    except Exception as e:
        logger.error(f"Error in Gmail auth callback: {e}")
        return f"Authorization failed: {str(e)}", 500

@app.route('/gmail/messages', methods=['GET'])
async def get_gmail_messages():
    """Get Gmail messages with optional filtering"""
    try:
        # Get query parameters
        folder = request.args.get('folder', 'INBOX')
        page = int(request.args.get('page', 1))
        page_size = int(request.args.get('pageSize', 20))
        unread_only = request.args.get('unread', 'false').lower() == 'true'
        search_query = request.args.get('q', '')
        
        # Get access token
        token = await get_gmail_access_token()
        
        if not token:
            return jsonify({
                'success': False,
                'error': 'Not authenticated with Gmail'
            })
        
        # Calculate pagination
        start_index = (page - 1) * page_size
        
        # Construct Gmail API query
        q = f"in:{folder}"
        
        if unread_only:
            q += " is:unread"
        
        if search_query:
            q += f" {search_query}"
        
        # Get messages list
        headers = {
            'Authorization': f'Bearer {token}'
        }
        
        # First get message IDs
        list_url = f"https://www.googleapis.com/gmail/v1/users/me/messages?q={urllib.parse.quote(q)}&maxResults={page_size}&startIndex={start_index}"
        response = requests.get(list_url, headers=headers)
        
        if response.status_code != 200:
            return jsonify({
                'success': False,
                'error': f"Failed to get messages: {response.text}"
            })
        
        data = response.json()
        messages = data.get('messages', [])
        total = data.get('resultSizeEstimate', 0)
        
        if not messages:
            return jsonify({
                'success': True,
                'messages': [],
                'total': 0,
                'unread': 0
            })
        
        # Get detailed message data
        detailed_messages = []
        
        for msg in messages:
            msg_id = msg['id']
            msg_url = f"https://www.googleapis.com/gmail/v1/users/me/messages/{msg_id}?format=metadata&metadataHeaders=From&metadataHeaders=To&metadataHeaders=Subject&metadataHeaders=Date"
            
            msg_response = requests.get(msg_url, headers=headers)
            if msg_response.status_code == 200:
                msg_data = msg_response.json()
                
                # Extract headers
                headers_dict = {}
                for header in msg_data.get('payload', {}).get('headers', []):
                    headers_dict[header['name']] = header['value']
                
                # Check if message is unread
                is_unread = 'UNREAD' in msg_data.get('labelIds', [])
                
                # Format message data
                detailed_messages.append({
                    'id': msg_id,
                    'threadId': msg_data.get('threadId', ''),
                    'from': headers_dict.get('From', 'Unknown Sender'),
                    'to': headers_dict.get('To', 'Unknown Recipient'),
                    'subject': headers_dict.get('Subject', '(No Subject)'),
                    'date': headers_dict.get('Date', ''),
                    'unread': is_unread,
                    'snippet': msg_data.get('snippet', '')
                })
        
        # Count unread messages
        unread_count = 0
        if folder == 'INBOX':
            unread_url = f"https://www.googleapis.com/gmail/v1/users/me/messages?q=in:inbox is:unread&maxResults=1"
            unread_response = requests.get(unread_url, headers=headers)
            
            if unread_response.status_code == 200:
                unread_data = unread_response.json()
                unread_count = unread_data.get('resultSizeEstimate', 0)
        
        return jsonify({
            'success': True,
            'messages': detailed_messages,
            'total': total,
            'unread': unread_count
        })
    except Exception as e:
        logger.error(f"Error getting Gmail messages: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        })

async def get_gmail_access_token():
    """Get a valid Gmail access token, refreshing if necessary"""
    try:
        async with aiosqlite.connect('db/assistant.db') as db:
            db.row_factory = sqlite3.Row
            
            # Check for existing token
            async with db.execute('SELECT * FROM gmail_auth_tokens ORDER BY id DESC LIMIT 1') as cursor:
                token_row = await cursor.fetchone()
                
                if not token_row:
                    return None
                
                token_data = dict(token_row)
                
                # Check if token is expired
                current_time = int(time.time())
                
                if current_time < token_data['expiry_time']:
                    # Token is still valid
                    return token_data['access_token']
                
                # Token is expired, try to refresh
                if not token_data['refresh_token']:
                    return None
                
                # Get Gmail API credentials from .env file
                client_id = os.getenv('GMAIL_CLIENT_ID', '')
                client_secret = os.getenv('GMAIL_CLIENT_SECRET', '')
                
                if not client_id or not client_secret:
                    return None
                
                # Refresh token
                refresh_data = {
                    'client_id': client_id,
                    'client_secret': client_secret,
                    'refresh_token': token_data['refresh_token'],
                    'grant_type': 'refresh_token'
                }
                
                token_response = requests.post(
                    'https://oauth2.googleapis.com/token',
                    data=refresh_data
                )
                
                if token_response.status_code != 200:
                    return None
                
                new_token_data = token_response.json()
                new_access_token = new_token_data.get('access_token')
                
                if not new_access_token:
                    return None
                
                # Update token in database
                new_expiry_time = int(time.time()) + new_token_data.get('expires_in', 3600)
                
                await db.execute('''
                    UPDATE gmail_auth_tokens 
                    SET access_token = ?, expiry_time = ?
                    WHERE id = ?
                ''', (new_access_token, new_expiry_time, token_data['id']))
                
                await db.commit()
                
                return new_access_token
    except Exception as e:
        logger.error(f"Error getting Gmail access token: {e}")
        return None

async def init_app():
    """Initialize the application"""
    try:
        global rag_system, calendar_rag, memory_system
        
        # Initialize database
        await init_db()
        
        # Initialize RAG system
        rag_system = EnhancedRAG()
        
        # Initialize calendar database
        await calendar_manager.setup_db()
        
        # Initialize memory system
        from memory import ContextualMemory
        memory_system = ContextualMemory(embeddings_model=embeddings_model)
        await memory_system.setup_db()
        
        # Sync RAG metadata with database
        await rag_system.sync_metadata_with_db()
        
        # Initialize Calendar RAG integration after rag_system is ready
        from calendar_integration import CalendarRAGIntegration
        calendar_rag = CalendarRAGIntegration(calendar_manager, rag_system)
        
        logger.info(f"Using Ollama model: {rag_system.model_name}")
    except Exception as e:
        logger.error(f"Error initializing app: {e}")
        logger.error(traceback.format_exc())

def open_animation():
    """Открывает анимацию в основном приложении"""
    import os
    
    # Получаем абсолютный путь к файлу анимации
    animation_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'wavy-animation.html')
    
    # Читаем содержимое файла анимации
    with open(animation_path, 'r') as f:
        animation_content = f.read()
    
    # Получаем абсолютный путь к файлу chat.html
    chat_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chat.html')
    
    # Читаем содержимое файла chat.html
    with open(chat_path, 'r') as f:
        chat_content = f.read()
    
    # Добавляем маршрут для анимации
    @app.route('/')
    async def index():
        return animation_content
    
    # Добавляем маршрут для chat.html
    @app.route('/chat.html')
    async def chat_page():
        return chat_content
    
    # Добавляем маршруты для JavaScript файлов
    @app.route('/app.js')
    async def app_js():
        return await send_file('app.js')
    
    @app.route('/utils.js')
    async def utils_js():
        return await send_file('utils.js')
    
    @app.route('/calendar.js')
    async def calendar_js():
        return await send_file('calendar.js')
    
    print("Анимация запуска добавлена в основное приложение")

# Memory management API
@app.route('/memory', methods=['GET', 'POST'])
async def memory_api():
    """API endpoint for memory management"""
    try:
        # Check if memory system is initialized
        if memory_system is None:
            return jsonify({
                'success': False,
                'error': 'Memory system is not initialized yet'
            }), 503  # Service Unavailable
            
        if request.method == 'GET':
            # Get memory stats
            stats = await memory_system.get_memory_stats()
            
            # Get memory settings
            async with aiosqlite.connect('db/assistant.db') as db:
                db.row_factory = sqlite3.Row
                async with db.execute('SELECT setting_name, setting_value FROM memory_settings') as cursor:
                    rows = await cursor.fetchall()
                    settings = {row['setting_name']: row['setting_value'] for row in rows}
            
            return jsonify({
                'success': True,
                'stats': stats,
                'settings': settings,
                'available_commands': await memory_system.get_memory_management_commands()
            })
        
        elif request.method == 'POST':
            data = await request.get_json()
            command = data.get('command')
            
            if not command:
                return jsonify({'error': 'No command provided'}), 400
            
            response = await memory_system.process_memory_command(command)
            
            return jsonify({
                'success': True,
                'response': response
            })
            
    except Exception as e:
        logger.error(f"Error in memory API: {str(e)}")
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500

# Main entry point
if __name__ == '__main__':
    # Добавляем анимацию в основное приложении
    open_animation()
    
    # Initialize the app before running
    loop = asyncio.get_event_loop()
    loop.run_until_complete(init_app())
    
    # Запускаем сервер
    app.run(host='0.0.0.0', port=5001, debug=True)